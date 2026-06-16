from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "DEVELOPER_MANUAL.md"
OUTPUT = ROOT / "artifacts" / "财务通用Agent-开发端手册.docx"

NAVY = "163A5F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GOLD = "B78628"
INK = "243342"
MUTED = "687786"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CODE_FILL = "F5F7FA"
NOTE_FILL = "FFF8E8"
WHITE = "FFFFFF"
RED = "9B1C1C"


def set_font(run, size=None, color=INK, bold=None, italic=None, ascii_font="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_table_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def create_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    level.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    return abstract_id


def create_numbering_instance(doc, abstract_id):
    numbering = doc.part.numbering_part.element
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.2


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        set_keep_table_row_together(row)
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, 9, MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    tail = paragraph.add_run(" 页")
    set_font(tail, 9, MUTED)


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_inline(paragraph, text, default_size=11, default_color=INK):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, default_size, default_color, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, default_size - 0.5, DARK_BLUE, ascii_font="Menlo", east_asia="Microsoft YaHei")
            run._element.get_or_add_rPr().append(_run_shading("EAF1F8"))
        else:
            run = paragraph.add_run(part)
            set_font(run, default_size, default_color)


def _run_shading(fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    return shd


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.36)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 9),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("财务通用 Agent  |  开发端手册")
    set_font(hr, 9, MUTED, bold=True)

    footer = section.footer
    add_page_number(footer.paragraphs[0])
    first_footer = section.first_page_footer
    fp = first_footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("finance-general-agent  |  V1.0  |  2026-06-06")
    set_font(fr, 9, MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FINANCE GENERAL AGENT")
    set_font(r, 11, GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("财务通用 Agent")
    set_font(r, 30, NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("开发端手册")
    set_font(r, 18, BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(48)
    r = p.add_run("本地启动 · 钉钉入口 · MCP联调 · Dubbo接入 · 冒烟测试")
    set_font(r, 11.5, MUTED)

    table = doc.add_table(rows=2, cols=3)
    set_table_geometry(table, [3000, 3000, 3000], indent_dxa=180)
    labels = [("76", "Gateway聚合"), ("45", "Finance MCP"), ("4", "Web页面"), ("22", "Browser MCP"), ("8", "Vision MCP"), ("1", "统一入口")]
    for idx, (metric, label) in enumerate(labels):
        cell = table.cell(idx // 3, idx % 3)
        set_cell_shading(cell, LIGHT_BLUE if idx < 3 else LIGHT_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(metric)
        set_font(r, 18, NAVY, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(label)
        set_font(r2, 9.5, MUTED, bold=True)
    set_repeat_table_header(table.rows[0])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(52)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("适用于后端开发、前端开发、MCP接入开发、测试与联调人员")
    set_font(r, 10.5, MUTED, italic=True)

    doc.add_page_break()


def add_quick_navigation(doc):
    p = doc.add_paragraph("快速导航", style="Heading 1")
    set_keep_with_next(p)
    intro = doc.add_paragraph()
    add_inline(intro, "首次联调建议依次阅读：**本地启动顺序 → 登录与入口说明 → MCP 与 Agent 调用方式 → 构建、测试与验收 → 常见排障**。")
    table = doc.add_table(rows=1, cols=3)
    headers = ["如果你想要", "查看章节", "关键命令或入口"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, 9.5, WHITE, bold=True)
    rows = [
        ("启动完整本地环境", "第 4 章", "按 Provider -> MCP -> Gateway -> Web 顺序启动"),
        ("模拟钉钉登录入口", "第 5 章", "http://localhost:5173/chat?channel=dingtalk..."),
        ("理解前后端模块", "第 6-8 章", "finance-web / finance-service / mcp"),
        ("通过 Gateway 调工具", "第 9 章", "POST http://localhost:9000/agent/chat"),
        ("新增真实 Dubbo 工具", "第 10 章", "API -> Core -> MCP -> Gateway -> 前端"),
        ("做联调验收", "第 12-13 章", "bash mcp/scripts/smoke-test.sh"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            add_inline(p, text, 9.5)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [2550, 1800, 4650])


def parse_markdown(doc, markdown):
    lines = markdown.splitlines()
    in_code = False
    code_lines = []
    table_lines = []
    skip_title_metadata = True
    decimal_abstract_id = create_decimal_numbering(doc)
    current_num_id = None
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        if skip_title_metadata:
            if stripped.startswith("## 1."):
                skip_title_metadata = False
            else:
                idx += 1
                continue

        if stripped.startswith("```"):
            current_num_id = None
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        if stripped.startswith("|"):
            current_num_id = None
            table_lines.append(stripped)
            idx += 1
            if idx >= len(lines) or not lines[idx].strip().startswith("|"):
                add_markdown_table(doc, table_lines)
                table_lines = []
            continue

        if not stripped:
            current_num_id = None
            idx += 1
            continue
        if stripped.startswith("### "):
            current_num_id = None
            doc.add_paragraph(stripped[4:], style="Heading 2")
        elif stripped.startswith("## "):
            current_num_id = None
            heading = stripped[3:]
            doc.add_paragraph(heading, style="Heading 1")
        elif stripped.startswith("> "):
            current_num_id = None
            add_note(doc, stripped[2:])
        elif re.match(r"^\d+\.\s+", stripped):
            if current_num_id is None:
                current_num_id = create_numbering_instance(doc, decimal_abstract_id)
            p = doc.add_paragraph()
            apply_numbering(p, current_num_id)
            add_inline(p, re.sub(r"^\d+\.\s+", "", stripped))
        elif stripped.startswith("- "):
            current_num_id = None
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
        else:
            current_num_id = None
            p = doc.add_paragraph()
            add_inline(p, stripped)
        idx += 1


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    set_paragraph_shading(p, NOTE_FILL)
    prefix = p.add_run("注意  ")
    set_font(prefix, 10.5, GOLD, bold=True)
    add_inline(p, text, 10.5, INK)


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.05
    set_paragraph_shading(p, CODE_FILL)
    for i, line in enumerate(lines):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        set_font(r, 8.4, DARK_BLUE, ascii_font="Menlo", east_asia="Microsoft YaHei")


def add_markdown_table(doc, lines):
    if len(lines) < 2:
        return
    rows = [[cell.strip() for cell in line.strip("|").split("|")] for line in lines]
    if all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=cols)
    header = rows[0]
    for i in range(cols):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header[i] if i < len(header) else "")
        set_font(r, 9.2, WHITE, bold=True)
    for row in rows[1:]:
        cells = table.add_row().cells
        for i in range(cols):
            p = cells[i].paragraphs[0]
            text = row[i] if i < len(row) else ""
            add_inline(p, text, 9.1)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == 1 and re.fullmatch(r"\d+", text):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    widths = table_widths(cols, rows)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)


def table_widths(cols, rows):
    if cols == 2:
        return [2700, 6300]
    if cols == 3:
        if any(row and row[0] == "服务" for row in rows):
            return [3100, 1500, 4400]
        return [2400, 2000, 4600]
    if cols == 4:
        return [1900, 1200, 2500, 3400]
    base = 9000 // cols
    return [base] * (cols - 1) + [9000 - base * (cols - 1)]


def add_end_page(doc):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(150)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("开始使用财务通用 Agent")
    set_font(r, 22, NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("从一句清晰、包含日期和业务对象的自然语言指令开始。")
    set_font(r, 11.5, MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("http://localhost:5173/chat")
    set_font(r, 12, BLUE, bold=True)


def main():
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_quick_navigation(doc)
    parse_markdown(doc, markdown)
    doc.core_properties.title = "财务通用 Agent 开发端手册"
    doc.core_properties.subject = "启动、登录、模块、联调与排障指南"
    doc.core_properties.author = "Finance General Agent"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
